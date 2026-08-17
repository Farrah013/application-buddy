"""Render an ATS-readable, review-only résumé DOCX."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def render_resume_docx(draft: dict[str, Any], output_path: str) -> str:
    """Use paragraphs only. ATS structure takes priority over visual styling."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6); section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7); section.right_margin = Inches(0.7)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"; styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].paragraph_format.space_after = Pt(2)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(draft["candidate"]["name"]); run.bold = True; run.font.name = "Arial"; run.font.size = Pt(18); run.font.color.rgb = RGBColor(23, 37, 84)
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(" | ".join([
        draft["candidate"]["location"],
        draft["candidate"]["phone"],
        draft["candidate"]["email"],
        draft["candidate"]["linkedin"],
        draft["candidate"]["portfolio"],
    ]))
    flag = doc.add_paragraph(); flag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    marker = flag.add_run("REVIEW DRAFT - NOT APPROVED FOR SUBMISSION"); marker.bold = True; marker.font.size = Pt(8); marker.font.color.rgb = RGBColor(153, 27, 27)

    def heading(text: str) -> None:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text.upper()); r.bold = True; r.font.name = "Arial"; r.font.size = Pt(10.5); r.font.color.rgb = RGBColor(23, 37, 84)

    heading("Professional Summary")
    doc.add_paragraph(draft["summary"]["text"])
    heading("Professional Experience")
    for role in draft["experience"]:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"{role['title']} | {role['employer']}"); r.bold = True
        p.add_run(f" | {role['location']} | {role['start']} - {role['end']}")
        for bullet in role["bullets"][:4]:
            bp = doc.add_paragraph(style="List Bullet"); bp.paragraph_format.left_indent = Inches(0.18); bp.paragraph_format.first_line_indent = Inches(-0.12); bp.paragraph_format.space_after = Pt(1)
            bp.add_run(bullet["text"])
    heading("Education")
    for item in draft["education"]:
        doc.add_paragraph(f"{item['degree']} | {item['institution']} | Expected {item['expected_graduation']}")
    if draft["skills"]:
        heading("Relevant Skills")
        doc.add_paragraph(" | ".join(draft["skills"]))
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
