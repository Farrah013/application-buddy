"""ATS-first résumé gate for Application Buddy."""

from __future__ import annotations

from pathlib import Path
from typing import Any
from zipfile import ZipFile

from docx import Document


REQUIRED_SECTIONS = ("PROFESSIONAL SUMMARY", "PROFESSIONAL EXPERIENCE", "EDUCATION")
FORBIDDEN_OBJECTS = ("w:txbxContent", "w:drawing", "w:pict")


def validate_resume_draft(draft: dict[str, Any]) -> dict[str, Any]:
    failures = []
    if draft.get("document_status") != "REVIEW_DRAFT":
        failures.append("Draft status is not REVIEW_DRAFT")
    if draft.get("final_use_allowed") is not False:
        failures.append("Final use is not blocked")
    if not draft.get("candidate", {}).get("name"):
        failures.append("Candidate name is missing")
    if not draft.get("candidate", {}).get("linkedin"):
        failures.append("Candidate LinkedIn URL is missing")
    if not draft.get("experience"):
        failures.append("Professional experience is missing")
    if not draft.get("education"):
        failures.append("Education is missing")
    for role in draft.get("experience", []):
        if not all(role.get(field) for field in ("employer", "title", "start", "end")):
            failures.append(f"Incomplete role record: {role.get('record_id', 'unknown')}")
        for bullet in role.get("bullets", []):
            if bullet.get("approval_status") != "SOURCE_VERIFIED":
                failures.append(f"Unapproved bullet: {bullet.get('id', 'unknown')}")
    return {"gate": "ATS_CONTENT", "passed": not failures, "failures": failures}


def validate_docx_structure(path: str) -> dict[str, Any]:
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    failures = []
    for section in REQUIRED_SECTIONS:
        if section not in text.upper():
            failures.append(f"Missing standard section: {section}")
    if document.tables:
        failures.append("Tables are not permitted in the ATS résumé template")
    with ZipFile(Path(path)) as archive:
        xml = b"\n".join(
            archive.read(name)
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )
    for marker in FORBIDDEN_OBJECTS:
        if marker.encode("utf-8") in xml:
            failures.append(f"Unsupported layout object detected: {marker}")
    if len(text.strip()) < 300:
        failures.append("Extracted résumé text is unexpectedly short")
    linkedin = "https://www.linkedin.com/in/farrah-j-8b6407176/"
    if linkedin not in text:
        failures.append("Exact LinkedIn URL is missing from selectable document text")
    return {
        "gate": "ATS_DOCUMENT",
        "passed": not failures,
        "failures": failures,
        "extracted_character_count": len(text.strip()),
        "required_sections": list(REQUIRED_SECTIONS),
    }


def run_ats_gate(draft: dict[str, Any], docx_path: str) -> dict[str, Any]:
    content = validate_resume_draft(draft)
    document = validate_docx_structure(docx_path)
    return {
        "priority": "ATS_FIRST",
        "passed": content["passed"] and document["passed"],
        "content": content,
        "document": document,
        "appearance_review_allowed": content["passed"] and document["passed"],
    }
